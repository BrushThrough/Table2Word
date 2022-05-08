import os
import cv2
import numpy as np
import paddlehub as hub
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor

def shrink_point(y,x):
    bitwise_and[y][x] = 0
    for i in range(8):
        tmp_y=int(dir[i][0])+y
        tmp_x=int(dir[i][1])+x
        if(bitwise_and[tmp_y][tmp_x]):
            shrink_point(tmp_y,tmp_x)

def judge_ver_exist_in_binary(y,x):
    if binary[y][x] :
        return 1
    for i in range(8):
        tmp_y=int(dir[i][0])+y
        tmp_x=int(dir[i][1])+x
        if binary[tmp_y][tmp_x]:
            return 1
    return 0

def judge_ver_exist_in_bitwise_and(y,x):
    if bitwise_and[y][x] :
        return 1
    for i in range(8):
        tmp_y=int(dir[i][0])+y
        tmp_x=int(dir[i][1])+x
        if bitwise_and[tmp_y][tmp_x]:
            return 1
    return 0

def judge_vertical_line(x,y_min,y_max):
    for i in range(10):
        if(judge_ver_exist_in_binary(y_min+(y_max-y_min)//10*i,x) == 0):
            return 0
    return 1

def judge_horizontal_line(y,x_min,x_max):
    for i in range(10):
        if(judge_ver_exist_in_binary(y,x_min+(x_max-x_min)//10*i) == 0):
            return 0
    return 1

def parse_pic_to_excel_data(src):
    raw = cv2.imread(src, 1)

    # 灰度图片
    gray = cv2.cvtColor(raw, cv2.COLOR_BGR2GRAY)

    # 二值化
    global binary
    binary = cv2.adaptiveThreshold(~gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 35, -5)
    # cv2.imshow("binary_picture", binary)  # 展示图片
    # cv2.waitKey()

    # 自适应获取核值 识别横线
    rows, cols = binary.shape
    scale = 30
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (cols // scale, 1))
    eroded = cv2.erode(binary, kernel, iterations=1)
    new_scale = 30
    new_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (cols // new_scale, 1))
    dilated_col = cv2.dilate(eroded, new_kernel, iterations=1)
    # cv2.imshow("excel_horizontal_line", dilated_col)
    # cv2.waitKey(0)
    # cv2.imwrite("col.jpg",dilated_col)

    # 识别竖线
    scale = 40
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, rows // scale))
    eroded = cv2.erode(binary, kernel, iterations=1)
    new_scale = 30
    new_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, rows // new_scale))
    dilated_row = cv2.dilate(eroded, new_kernel, iterations=1)
    # cv2.imshow("excel_vertical_line", dilated_row)
    # cv2.waitKey(0)
    # cv2.imwrite("row.jpg",dilated_row)

    # 标识交点
    global bitwise_and
    bitwise_and = cv2.bitwise_and(dilated_col, dilated_row)
    # cv2.imshow("excel_bitwise_and", bitwise_and)
    # cv2.waitKey(0)
    # cv2.imwrite("bitwise_and.jpg",bitwise_and)

    # 标识表格
    merge = cv2.add(dilated_col, dilated_row)
    # cv2.imshow("entire_excel_contour", merge)merge
    # cv2.waitKey(0)
    # cv2.imwrite("merge.jpg",)

    # 识别黑白图中的白色交叉点，将横纵坐标取出
    ys, xs = np.where(bitwise_and > 0)

    # 缩点
    global dir
    dir = np.zeros((8,2))
    dir[0][0]=1;dir[1][0]=-1;dir[2][1]=1;dir[3][1]=-1
    dir[4][0]=1;dir[4][1]=1
    dir[5][0]=1;dir[5][1]=-1
    dir[6][0]=-1;dir[6][1]=1
    dir[7][0]=-1;dir[7][1]=-1


    for i in range(len(xs)):
        if(bitwise_and[ys[i]][xs[i]]):
            shrink_point(ys[i], xs[i])
            bitwise_and[ys[i]][xs[i]]=255


    # cv2.imwrite("bitwise_and.jpg", bitwise_and)
    # cv2.waitKey(0)

    # 识别黑白图(缩点后)中的白色交叉点，将横纵坐标取出
    ys, xs = np.where(bitwise_and > 0)
    ys, xs = np.where(bitwise_and > 0)
    # 纵坐标
    global y_point_arr
    y_point_arr = []
    # 横坐标
    global x_point_arr
    x_point_arr = []

    # 对X和Y坐标排序和判重
    i = 0
    sort_x_point = np.sort(xs)
    for i in range(len(sort_x_point) - 1):
        if sort_x_point[i + 1] - sort_x_point[i] > 1:
            x_point_arr.append(sort_x_point[i])
    x_point_arr.append(sort_x_point[i])

    i = 0
    sort_y_point = np.sort(ys)
    # print(np.sort(ys))
    for i in range(len(sort_y_point) - 1):
        if sort_y_point[i + 1] - sort_y_point[i] > 1:
            y_point_arr.append(sort_y_point[i])
    y_point_arr.append(sort_y_point[i])

    global is_table_vertex
    global table_vertex_horizontal_num
    global table_vertex_vertical_num
    is_table_vertex = [[0]*len(x_point_arr) for _ in range(len(y_point_arr))]
    table_vertex_horizontal_num = [[0]*len(x_point_arr) for _ in range(len(y_point_arr))]
    table_vertex_vertical_num = [[0]*len(x_point_arr) for _ in range(len(y_point_arr))]

    # print('y_point_arr', y_point_arr)
    # print('x_point_arr', x_point_arr)

    global num
    num=0
    os.chdir("seg_images")

    # 循环y坐标，x坐标分割表格
    for i in range(len(y_point_arr) - 1):
        for j in range(len(x_point_arr) - 1):
            if judge_ver_exist_in_bitwise_and(y_point_arr[i],x_point_arr[j]) ==0:
                continue
            aim_i=0
            aim_j=0
            for ti in range(i+1,len(y_point_arr)):
                for tj in range(j+1,len(x_point_arr)):
                    if judge_ver_exist_in_bitwise_and(y_point_arr[ti],x_point_arr[tj]) == 0:
                        continue
                    if(judge_horizontal_line(y_point_arr[i],x_point_arr[j],x_point_arr[tj]) and
                       judge_horizontal_line(y_point_arr[ti],x_point_arr[j],x_point_arr[tj]) and
                       judge_vertical_line(x_point_arr[j],y_point_arr[i],y_point_arr[ti]) and
                       judge_vertical_line(x_point_arr[tj], y_point_arr[i], y_point_arr[ti]) and
                       (y_point_arr[ti] - y_point_arr[i])>10 and
                       (x_point_arr[tj] - x_point_arr[j])>10):
                            aim_i = ti
                            aim_j = tj
                            break
                if aim_i:
                    break
            if aim_i==0:
                continue
            is_table_vertex[i][j]=1
            table_vertex_vertical_num[i][j] = aim_i - i
            table_vertex_horizontal_num[i][j] = aim_j -j
            # print(y_point_arr[i],x_point_arr[j])
            # print(y_point_arr[aim_i],x_point_arr[aim_j])
            # 在分割时，第一个参数为y坐标，第二个参数为x坐标
            cell = raw[y_point_arr[i]:y_point_arr[aim_i], x_point_arr[j]:x_point_arr[aim_j]]
            cell_2 = cv2.resize(cell, None, fx=4, fy=4, interpolation=cv2.INTER_LINEAR)
            # cv2.imshow("seg",cell)
            # cv2.imshow("seg_2", cell_2)
            # cv2.waitKey(0)
            cv2.imwrite(str(num)+".jpg",cell)
            num += 1


def text_detection_recognition():
    os.chdir("..")
    # 加载移动端预训练模型
    ocr = hub.Module(name="chinese_ocr_db_crnn_mobile")
    # 服务端可以加载大模型，效果更好
    # ocr = hub.Module(name="chinese_ocr_db_crnn_server")

    test_img_path = []
    for i in range(num):
        test_img_path.append("seg_images/"+str(i) + ".jpg")

    # 读取测试文件夹test.txt中的照片路径
    np_images = [cv2.imread(image_path) for image_path in test_img_path]

    global results
    results = ocr.recognize_text(
        images=np_images,  # 图片数据，ndarray.shape 为 [H, W, C]，BGR格式；
        use_gpu=False,  # 是否使用 GPU；若使用GPU，请先设置CUDA_VISIBLE_DEVICES环境变量
        output_dir='ocr_result',  # 图片的保存路径，默认设为 ocr_result；
        visualization=True,  # 是否将识别结果保存为图片文件；
        box_thresh=0.5,  # 检测文本框置信度的阈值；
        text_thresh=0.5)  # 识别中文文本置信度的阈值；
'''
    for result in results:
        data = result['data']
        save_path = result['save_path']
        for infomation in data:
             print('text: ', infomation['text'], '\nconfidence: ', infomation['confidence'], '\ntext_box_position: ', infomation['text_box_position'])
'''

def table_restore():
    document = Document()
    document.styles['Normal'].font.name = u'仿宋'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    row_num = len(y_point_arr)-1
    col_num = len(x_point_arr)-1
    table = document.add_table(rows = row_num, cols = 0, style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j in range(col_num):
        table.add_column(width=Cm(16.5*(x_point_arr[j+1] - x_point_arr[j])/(x_point_arr[-1] - x_point_arr[0])))

    for i in range(row_num):
        table.rows[i].height = Cm(20.0*(y_point_arr[i+1] - y_point_arr[i]-1)/(y_point_arr[-1] - y_point_arr[0]))

    for i,row in enumerate(table.rows):
        for j,cell in enumerate(row.cells):
            if is_table_vertex[i][j] == 0:
                continue
            for ti in range(table_vertex_vertical_num[i][j]):
                for tj in range(table_vertex_horizontal_num[i][j]):
                        table.cell(i+ti, j).merge(table.cell(i+ti, j+tj))

            for ti in range(table_vertex_vertical_num[i][j]):
                table.cell(i, j).merge(table.cell(i + ti, j))
    pointer = 0
    for i,row in enumerate(table.rows):
        for j,cell in enumerate(row.cells):
            if is_table_vertex[i][j] == 0:
                continue
            result  = results[pointer]
            pointer += 1
            data = result['data']
            last_y = 0
            for infomation in data:
                if table.cell(i, j).text != "" and abs(infomation['text_box_position'][0][1] - last_y)>=10:
                    table.cell(i, j).text += "\n"

                table.cell(i,j).text+=infomation['text']
                last_y = infomation['text_box_position'][0][1]
            table.cell(i,j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直居中
            table.cell(i,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 水平居中
    document.save('result.docx')

if __name__ == '__main__':
    file = "test_imgs/1.png"
    parse_pic_to_excel_data(file)
    text_detection_recognition()
    table_restore()