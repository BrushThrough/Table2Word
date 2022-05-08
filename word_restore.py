# coding:utf-8
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
import os

class Word_Maker():
    def __init__(self):
        pass
    def __call__(self,x_point_arr ,y_point_arr,is_table_vertex,table_vertex_horizontal_num, table_vertex_vertical_num,ocr_results,dir):

        document = Document()

        #word格式设置
        document.styles['Normal'].font.name = u'仿宋'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        document.styles['Normal'].font.size = Pt(10.5)
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # 还原冗余表格框架
        row_num = len(y_point_arr) - 1
        col_num = len(x_point_arr) - 1
        table = document.add_table(rows=row_num, cols=0, style='Table Grid')
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 还原冗余表格大小
        for j in range(col_num):
            table.add_column(
                width=Cm(16.5 * (x_point_arr[j + 1] - x_point_arr[j]) / (x_point_arr[-1] - x_point_arr[0])))

        for i in range(row_num):
            table.rows[i].height = Cm(
                20.0 * (y_point_arr[i + 1] - y_point_arr[i] - 1) / (y_point_arr[-1] - y_point_arr[0]))

        #合并冗余单元格
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if is_table_vertex[i][j] == 0:
                    continue
                for ti in range(table_vertex_vertical_num[i][j]):
                    for tj in range(table_vertex_horizontal_num[i][j]):
                        table.cell(i + ti, j).merge(table.cell(i + ti, j + tj))

                for ti in range(table_vertex_vertical_num[i][j]):
                    table.cell(i, j).merge(table.cell(i + ti, j))

        # 还原表格文字
        pointer = 0
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if is_table_vertex[i][j] == 0:
                    continue
                ocr_result = ocr_results[pointer]
                pointer += 1
                data = ocr_result['data']
                last_y = 0
                for infomation in data:
                    if table.cell(i, j).text != "" and abs(infomation['text_box_position'][0][1] - last_y) >= 10:
                        table.cell(i, j).text += "\n"

                    table.cell(i, j).text += infomation['text']
                    last_y = infomation['text_box_position'][0][1]
                table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直居中
                table.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 水平居中
        document.save(dir+'_result.docx')
